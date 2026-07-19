# -*- coding: utf-8 -*-
"""
my-toolbox 全功能自动化测试程序
================================

对在线工具箱做端到端冒烟测试：
  1. 应用能否正常启动（捕获所有工具模块的导入/语法错误）
  2. 首页是否渲染并列出全部已启用工具
  3. 每个工具的页面（GET）是否正常
  4. 每个工具的核心处理逻辑（POST）是否产出正确结果
  5. 鉴权流程：注册 / 登录 / 登出
  6. 管理后台：仪表盘 / 用户 / 工具 / 日志 / 设置 / 各项操作
  7. 匿名试用配额是否生效

用法：
    python tests/run_tests.py

输出：
    - 控制台 PASS/FAIL 汇总
    - tests/test_report.json  机器可读结果
    - tests/test_report.html  可视化报告

注意：本脚本使用内存数据库与临时目录，不会污染项目数据。
"""
from __future__ import annotations

import io
import json
import os
import sys
import tempfile
import time
import traceback
from datetime import datetime, timezone
from pathlib import Path

# ---------------------------------------------------------------------------
# 0. 测试环境配置（必须在 import app 之前设置）
# ---------------------------------------------------------------------------
PROJECT_DIR = Path(__file__).resolve().parent.parent
os.chdir(PROJECT_DIR)
# 确保项目根目录在导入路径中（app / tools 等模块才能被 import）
if str(PROJECT_DIR) not in sys.path:
    sys.path.insert(0, str(PROJECT_DIR))

# 模块化导入失败记录（在探查阶段填充）
failed_imports: dict = {}


def _write_reports():
    """生成本地 JSON / HTML 测试报告。"""
    total = len(RESULTS)
    passed = sum(1 for r in RESULTS if r["ok"])
    failed = total - passed

    out_dir = PROJECT_DIR / "tests"
    out_dir.mkdir(exist_ok=True)

    # JSON
    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "summary": {"total": total, "passed": passed, "failed": failed},
        "failed_tool_imports": failed_imports,
        "results": RESULTS,
    }
    out_json = out_dir / "test_report.json"
    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    # HTML
    failed_rows = "".join(
        f"<tr class='fail'><td>{r['group']}</td><td>{r['name']}</td><td>{r['message']}</td></tr>"
        for r in RESULTS if not r["ok"]
    ) or "<tr><td colspan='3'>无失败项 🎉</td></tr>"
    rows = "".join(
        f"<tr class='{'pass' if r['ok'] else 'fail'}'><td>{r['group']}</td><td>{r['name']}</td>"
        f"<td>{'✅' if r['ok'] else '❌'}</td><td>{r['message']}</td></tr>"
        for r in RESULTS
    )
    html = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>my-toolbox 测试报告</title>
<style>
 body{{font-family:system-ui,'Segoe UI',sans-serif;margin:24px;color:#222}}
 h1{{margin-bottom:4px}} .summary{{margin:12px 0;font-size:15px}}
 .badge{{display:inline-block;padding:4px 10px;border-radius:8px;color:#fff;margin-right:8px}}
 .b-pass{{background:#198754}} .b-fail{{background:#dc3545}}
 table{{border-collapse:collapse;width:100%;margin-top:16px;font-size:13px}}
 th,td{{border:1px solid #ddd;padding:6px 8px;text-align:left;vertical-align:top}}
 tr.pass td{{background:#f4fcf6}} tr.fail td{{background:#fdf4f4}}
 th{{background:#f0f0f0}}
 h2{{margin-top:28px}}
 code{{background:#f3f3f3;padding:1px 4px;border-radius:4px}}
</style></head><body>
<h1>my-toolbox 全功能测试报告</h1>
<div class="summary">
 <span class="badge b-pass">通过 {passed}</span>
 <span class="badge b-fail">失败 {failed}</span>
 <span>总计 {total}</span> · 生成时间 {report['generated_at']}
</div>
<h2>失败项明细</h2>
<table><thead><tr><th>分组</th><th>检查项</th><th>说明</th></tr></thead>
<tbody>{failed_rows}</tbody></table>
<h2>全部检查</h2>
<table><thead><tr><th>分组</th><th>检查项</th><th>结果</th><th>说明</th></tr></thead>
<tbody>{rows}</tbody></table>
</body></html>"""
    out_html = out_dir / "test_report.html"
    out_html.write_text(html, encoding="utf-8")

    print("\n" + "=" * 72)
    print(f"测试完成：共 {total} 项，通过 {passed}，失败 {failed}")
    print(f"报告已生成：{out_json}  /  {out_html}")
    print("=" * 72)
    return out_json, out_html

# 使用内存数据库，避免污染本地 app.db
os.environ.setdefault("DATABASE_URL", "sqlite:///:memory:")
# 开发模式：关闭 secure cookie，便于测试客户端携带 session
os.environ["FLASK_ENV"] = "development"
# 关闭后台调度（避免测试进程挂起 / 拉起线程）
os.environ["VERCEL"] = "1"  # 让 create_app 跳过 APScheduler 分支
# AI 作图使用 mock 提供方，避免联网依赖
os.environ["AI_PROVIDER"] = "mock"
# 固定的种子管理员账号（使用 .com 域名，email-validator 会拒绝 .local）
os.environ["ADMIN_EMAIL"] = "admin@test.com"
os.environ["ADMIN_PASSWORD"] = "Admin123456"
os.environ["SECRET_KEY"] = "test-secret-key-not-for-prod"

# 结果收集
RESULTS: list[dict] = []


def record(group: str, name: str, ok: bool, message: str = "", detail: str = "") -> None:
    RESULTS.append({
        "group": group,
        "name": name,
        "ok": bool(ok),
        "message": message,
        "detail": detail,
    })
    status = "PASS" if ok else "FAIL"
    line = f"[{status}] {group} :: {name}"
    if message:
        line += f" — {message}"
    print(line, flush=True)


# 临时目录（上传 / 实例）
_TMP = tempfile.mkdtemp(prefix="mytoolbox_test_")

# ---------------------------------------------------------------------------
# 1. 启动应用（捕获导入错误）
# ---------------------------------------------------------------------------
print("=" * 72)
print("启动 my-toolbox 测试环境 ...")
print("=" * 72)

try:
    from app import create_app
    app = create_app()
    record("启动", "应用工厂 create_app()", True, "应用启动成功")
except Exception as exc:  # noqa: BLE001
    record("启动", "应用工厂 create_app()", False, f"启动失败: {exc}", traceback.format_exc())
    # 即便启动失败，也尝试单独导入各工具模块以定位问题
    app = None

# 即便整体启动失败，也逐个探测工具模块，定位是哪个模块拖垮了启动
import importlib

TOOL_IDS = []
try:
    import yaml
    cfg = yaml.safe_load((PROJECT_DIR / "tools_config.yaml").read_text(encoding="utf-8")) or {}
    TOOL_IDS = [t["id"] for t in cfg.get("tools", [])]
except Exception as exc:  # noqa: BLE001
    record("启动", "读取 tools_config.yaml", False, str(exc))

for tid in TOOL_IDS:
    mod = f"tools.{tid}"
    try:
        importlib.import_module(mod)
        record("工具模块", f"import {mod}", True)
    except Exception as exc:  # noqa: BLE001
        failed_imports[tid] = repr(exc)
        record("工具模块", f"import {mod}", False, str(exc).splitlines()[0], traceback.format_exc())

if app is None:
    # 无法继续做路由级测试
    _write_reports()
    print("\n应用无法启动，停止后续测试。请先修复上述导入错误。")
    sys.exit(1)

# ---------------------------------------------------------------------------
# 2. 测试客户端 & 运行时配置
# ---------------------------------------------------------------------------
app.config["WTF_CSRF_ENABLED"] = False  # 关闭 CSRF 以便测试提交表单
app.config["TESTING"] = True
app.config["UPLOAD_DIR"] = Path(_TMP) / "uploads"
app.config["INSTANCE_DIR"] = Path(_TMP) / "instance"
app.config["UPLOAD_DIR"].mkdir(parents=True, exist_ok=True)
app.config["INSTANCE_DIR"].mkdir(parents=True, exist_ok=True)

client = app.test_client()

# 健康检查
try:
    r = client.get("/healthz")
    ok = r.status_code == 200 and r.get_json().get("status") == "ok"
    record("系统", "GET /healthz", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("系统", "GET /healthz", False, str(exc))

# 首页
try:
    r = client.get("/")
    body = r.get_data(as_text=True)
    enabled = [t.id for t in app.db.session.query(__import__("models").Tool).filter_by(enabled=True).all()] \
        if False else None
    # 简单判断：首页 200 且包含若干工具名称
    ok = r.status_code == 200 and ("PDF" in body or "工具" in body)
    record("系统", "GET / 首页渲染", ok, f"status={r.status_code}, bytes={len(body)}")
except Exception as exc:  # noqa: BLE001
    record("系统", "GET / 首页渲染", False, str(exc))


# ---------------------------------------------------------------------------
# 3. 测试数据夹具生成
# ---------------------------------------------------------------------------
def make_pdf(n_pages: int = 1) -> io.BytesIO:
    from pypdf import PdfWriter
    buf = io.BytesIO()
    w = PdfWriter()
    for _ in range(n_pages):
        w.add_blank_page(width=200, height=200)
    w.write(buf)
    buf.seek(0)
    return buf


def make_png(w=64, h=64, color=(255, 0, 0)) -> io.BytesIO:
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (w, h), color).save(buf, format="PNG")
    buf.seek(0)
    return buf


def make_docx(text: str = "Hello world") -> io.BytesIO:
    from docx import Document
    d = Document()
    d.add_heading("测试标题", level=1)
    d.add_paragraph(text)
    d.add_paragraph("第二段内容用于验证多段落转换。")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


def make_fcst_xlsx() -> io.BytesIO:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    months = ["2024-01", "2024-02", "2024-03", "2024-04", "2024-05", "2024-06"]
    for i, m in enumerate(months):
        ws.cell(row=1, column=10 + i, value=m)  # J..O 列放月份表头
    data = [
        ("PN001", [10, 20, 30, 40, 50, 60]),
        ("PN002", [1, 2, 3, 4, 5, 6]),
        ("PN003", [7, 8, 9, 10, 11, 12]),
    ]
    for r, (pn, vals) in enumerate(data):
        excel_row = 2 + r * 4  # 第 2/6/10 行（每 4 行一个品号）
        ws.cell(row=excel_row, column=5, value=pn)  # E 列 = 品号
        for i, v in enumerate(vals):
            ws.cell(row=excel_row, column=10 + i, value=v)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


AJ = {"Accept": "application/json"}  # 让工具返回 JSON 而非跳转


def _normalize_form(data):
    """将 (key, value) 列表规整为 dict；同名多值（如多文件）归并为列表。"""
    if isinstance(data, dict):
        return data
    d: dict = {}
    for k, v in data:
        if k in d:
            if not isinstance(d[k], list):
                d[k] = [d[k]]
            d[k].append(v)
        else:
            d[k] = v
    return d


def post_tool(tool_id: str, data, headers=None):
    """向 /tools/<route-base>/process 发起 POST。route base 从 yaml 推断。"""
    import yaml
    cfg = yaml.safe_load((PROJECT_DIR / "tools_config.yaml").read_text(encoding="utf-8")) or {}
    route = next((t["route"] for t in cfg.get("tools", []) if t["id"] == tool_id), f"/tools/{tool_id}")
    url = f"{route}/process"
    h = dict(AJ)
    if headers:
        h.update(headers)
    return client.post(url, data=_normalize_form(data), content_type="multipart/form-data", headers=h)


# ---------------------------------------------------------------------------
# 4. 每个工具的页面 GET
# ---------------------------------------------------------------------------
import yaml as _yaml
_cfg = _yaml.safe_load((PROJECT_DIR / "tools_config.yaml").read_text(encoding="utf-8")) or {}
for t in _cfg.get("tools", []):
    tid = t["id"]
    route = t["route"]
    try:
        # strict_slashes=False 后，无尾斜杠也应直接 200（不再 308 重定向）
        r1 = client.get(route)
        r2 = client.get(route + "/")
        body = r2.get_data(as_text=True)
        has_content = (f"{route}/process" in body) or (f"{route}/generate" in body) or ("<form" in body) or ("<button" in body)
        ok = (r1.status_code == 200 and r2.status_code == 200 and has_content)
        record("工具页面", f"GET {route} ({tid})", ok,
               f"no-slash={r1.status_code} w-slash={r2.status_code} body={has_content}")
    except Exception as exc:  # noqa: BLE001
        record("工具页面", f"GET {route} ({tid})", False, str(exc).splitlines()[0], traceback.format_exc())


# ---------------------------------------------------------------------------
# 5. 每个工具的功能测试（POST）
# ---------------------------------------------------------------------------
def _assert_json_ok(r, extra=None):
    if r.status_code != 200:
        return False, f"HTTP {r.status_code}: {r.get_data(as_text=True)[:200]}"
    j = r.get_json(silent=True)
    if not isinstance(j, dict):
        return False, f"响应不是 JSON: {r.get_data(as_text=True)[:200]}"
    if not j.get("ok"):
        return False, f"ok=False: {j.get('error', j)}"
    if extra:
        for k in extra:
            if k not in j:
                return False, f"响应缺少字段 {k}: {j}"
    return True, ""


# 5.1 PDF 合并
try:
    data = [
        ("pdfs", (make_pdf(2), "a.pdf")),
        ("pdfs", (make_pdf(1), "b.pdf")),
    ]
    r = post_tool("pdf_merge", data)
    ok, msg = _assert_json_ok(r, ["url", "size"])
    record("功能", "pdf_merge 合并两个 PDF", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "pdf_merge 合并两个 PDF", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.2 PDF 拆分
try:
    data = [("pdf", (make_pdf(3), "src.pdf")), ("ranges", "1-2")]
    r = post_tool("pdf_split", data)
    ok, msg = _assert_json_ok(r, ["size"])
    record("功能", "pdf_split 拆分 PDF (1-2)", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "pdf_split 拆分 PDF (1-2)", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.3 PDF 加水印
try:
    data = [("pdf", (make_pdf(2), "src.pdf")), ("text", "机密")]
    r = post_tool("pdf_watermark", data)
    ok, msg = _assert_json_ok(r, ["page_count"])
    record("功能", "pdf_watermark 添加水印", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "pdf_watermark 添加水印", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.4 PDF 旋转
try:
    data = [("pdf", (make_pdf(2), "src.pdf")), ("angle", "90"), ("scope", "all")]
    r = post_tool("pdf_rotate", data)
    ok, msg = _assert_json_ok(r, ["rotated_pages"])
    record("功能", "pdf_rotate 旋转 90°", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "pdf_rotate 旋转 90°", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.5 PDF 转 Word
try:
    data = [("pdf", (make_pdf(2), "src.pdf"))]
    r = post_tool("pdf_to_word", data)
    ok, msg = _assert_json_ok(r, ["size"])
    record("功能", "pdf_to_word PDF 转 Word", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "pdf_to_word PDF 转 Word", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.6 图片压缩
try:
    data = [("image", (make_png(), "pic.png")), ("quality", "60")]
    r = post_tool("image_compress", data)
    ok, msg = _assert_json_ok(r, ["size"])
    record("功能", "image_compress 压缩图片", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "image_compress 压缩图片", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.7 AI 作图（mock）— 该工具端点为 /generate 而非 /process
try:
    r = client.post("/tools/ai-image/generate",
                    data={"prompt": "a red cat", "size": "512x512"},
                    content_type="multipart/form-data", headers=AJ)
    if r.status_code == 200 and (j := r.get_json(silent=True)) and j.get("task_id"):
        tid = j["task_id"]
        s = client.get(f"/tools/ai-image/status/{tid}", headers=AJ)
        sj = s.get_json(silent=True) or {}
        ok = sj.get("status") == "done" and sj.get("url")
        msg = "" if ok else f"status={sj.get('status')} err={sj.get('error')}"
        record("功能", "ai_image AI 作图 (mock)", ok, msg)
    else:
        record("功能", "ai_image AI 作图 (mock)", False, f"未返回 task_id: {r.get_data(as_text=True)[:120]}")
except Exception as exc:  # noqa: BLE001
    record("功能", "ai_image AI 作图 (mock)", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.8 FCST 合并 + PN 映射 CRUD
try:
    data = [("files", (make_fcst_xlsx(), "fcst.xlsx"))]
    r = post_tool("fcst_merge", data)
    ok, msg = _assert_json_ok(r, ["stats"])
    if ok:
        st = r.get_json()["stats"]
        ok = st.get("part_numbers", 0) == 3
        msg = f"品号数={st.get('part_numbers')}"
    record("功能", "fcst_merge FCST 预测合并", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "fcst_merge FCST 预测合并", False, str(exc).splitlines()[0], traceback.format_exc())

# FCST 需要登录（owner 隔离）。先登录管理员再做 PN CRUD 测试
def login_admin():
    # 先登出，确保从干净会话登录为管理员（已登录用户会被 login 直接重定向首页）
    client.get("/logout", follow_redirects=False)
    return client.post("/login", data={
        "email": os.environ["ADMIN_EMAIL"],
        "password": os.environ["ADMIN_PASSWORD"],
        "remember": "y",
    }, follow_redirects=False)


try:
    lr = login_admin()
    admin_logged_in = lr.status_code in (302, 302) or "home" in (lr.headers.get("Location", ""))
    record("鉴权", "管理员登录", lr.status_code in (301, 302), f"status={lr.status_code}")

    # PN 创建
    r = client.post("/tools/fcst-merge/api/pn",
                    json={"part_number": "PNX1", "mfr_part": "M-1", "brand": "B-1"},
                    headers={"Content-Type": "application/json"})
    ok = r.status_code == 201 and (r.get_json(silent=True) or {}).get("ok")
    record("功能", "fcst_merge PN 映射创建", ok, f"status={r.status_code}")

    # PN 列表
    r = client.get("/tools/fcst-merge/api/pn", headers=AJ)
    ok = (r.get_json(silent=True) or {}).get("ok") and (r.get_json(silent=True) or {}).get("total", 0) >= 1
    record("功能", "fcst_merge PN 映射列表", ok, f"status={r.status_code}")

    # PN 更新
    rid = (client.get("/tools/fcst-merge/api/pn", headers=AJ).get_json() or {}).get("items", [{}])[0].get("id")
    if rid:
        r = client.put(f"/tools/fcst-merge/api/pn/{rid}",
                       json={"brand": "B-2"}, headers={"Content-Type": "application/json"})
        ok = (r.get_json(silent=True) or {}).get("ok") and (r.get_json(silent=True) or {}).get("item", {}).get("brand") == "B-2"
        record("功能", "fcst_merge PN 映射更新", ok, f"status={r.status_code}")
    else:
        record("功能", "fcst_merge PN 映射更新", False, "未获取到记录 id")

    # PN 删除
    if rid:
        r = client.delete(f"/tools/fcst-merge/api/pn/{rid}", headers={"Content-Type": "application/json"})
        ok = (r.get_json(silent=True) or {}).get("ok")
        record("功能", "fcst_merge PN 映射删除", ok, f"status={r.status_code}")
    else:
        record("功能", "fcst_merge PN 映射删除", False, "未获取到记录 id")
except Exception as exc:  # noqa: BLE001
    record("功能", "fcst_merge PN CRUD", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.9 图片格式转换 PNG->JPG
try:
    data = [("file", (make_png(), "pic.png")), ("format", "jpg"), ("quality", "85")]
    r = post_tool("image_convert", data)
    ok, msg = _assert_json_ok(r, ["mime"])
    record("功能", "image_convert PNG→JPG", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "image_convert PNG→JPG", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.10 图片转 PDF
try:
    data = [("files", (make_png(), "pic.png"))]
    r = post_tool("image_to_pdf", data)
    ok, msg = _assert_json_ok(r, ["mime"])
    record("功能", "image_to_pdf 图片转 PDF", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "image_to_pdf 图片转 PDF", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.11 JSON 格式化
try:
    data = [("text", '{"b":2,"a":1}'), ("action", "format")]
    r = post_tool("json_formatter", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = "\n" in r.get_json()["result"]  # 美化后含换行
        msg = "结果已美化" if ok else "结果未美化"
    record("功能", "json_formatter 格式化", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "json_formatter 格式化", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.12 时间戳转换
try:
    data = [("direction", "ts2dt"), ("value", "1700000000")]
    r = post_tool("timestamp", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = "2023" in r.get_json()["result"]["utc"]
        msg = r.get_json()["result"]["utc"]
    record("功能", "timestamp 秒戳→日期", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "timestamp 秒戳→日期", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.13 Base64 编解码
try:
    data = [("action", "encode"), ("mode", "text"), ("text", "hello")]
    r = post_tool("base64", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        import base64 as _b64
        ok = _b64.b64decode(r.get_json()["result"]).decode() == "hello"
        msg = "编解码一致"
    record("功能", "base64 文本编码", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "base64 文本编码", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.14 URL 编解码
try:
    data = [("action", "encode"), ("text", "a b/c")]
    r = post_tool("url_codec", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = "%20" in r.get_json()["result"] or "%2F" in r.get_json()["result"]
        msg = r.get_json()["result"]
    record("功能", "url_codec 编码", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "url_codec 编码", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.15 UUID 生成
try:
    data = [("count", "5"), ("version", "4")]
    r = post_tool("uuid_gen", data)
    ok, msg = _assert_json_ok(r, ["uuids"])
    if ok:
        ok = len(r.get_json()["uuids"]) == 5
        msg = f"生成 {len(r.get_json()['uuids'])} 个"
    record("功能", "uuid_gen 生成 UUID", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "uuid_gen 生成 UUID", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.16 二维码生成
try:
    data = [("text", "https://example.com"), ("size", "10")]
    r = post_tool("qr_gen", data)
    ok, msg = _assert_json_ok(r, ["mime"])
    if ok:
        ok = r.get_json()["mime"] == "image/png"
    record("功能", "qr_gen 生成二维码", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "qr_gen 生成二维码", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.17 哈希计算
try:
    import hashlib as _hl
    expect = _hl.sha256(b"hello").hexdigest()
    data = [("algo", "sha256"), ("text", "hello")]
    r = post_tool("hash_calc", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = r.get_json()["result"] == expect
        msg = "sha256 一致" if ok else "sha256 不一致"
    record("功能", "hash_calc SHA256", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "hash_calc SHA256", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.18 正则测试器
try:
    data = [("pattern", r"\d+"), ("text", "abc123def456"), ("flags", "")]
    r = post_tool("regex_tester", data)
    ok, msg = _assert_json_ok(r, ["count"])
    if ok:
        ok = r.get_json()["count"] == 2
        msg = f"匹配 {r.get_json()['count']} 处"
    record("功能", "regex_tester 正则匹配", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "regex_tester 正则匹配", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.19 SQL 格式化
try:
    data = [("sql", "select * from users where id=1"), ("action", "format")]
    r = post_tool("sql_formatter", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = "SELECT" in r.get_json()["result"]
        msg = "已转大写"
    record("功能", "sql_formatter 格式化", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "sql_formatter 格式化", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.20 Word 转 PDF
try:
    data = [("file", (make_docx(), "doc.docx"))]
    r = post_tool("word_to_pdf", data)
    ok, msg = _assert_json_ok(r, ["mime"])
    if ok:
        ok = r.get_json()["mime"] == "application/pdf"
    record("功能", "word_to_pdf Word 转 PDF", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "word_to_pdf Word 转 PDF", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.21 Markdown 预览
try:
    data = [("text", "# 标题\n\n这是**加粗**内容。")]
    r = post_tool("markdown_preview", data)
    ok, msg = _assert_json_ok(r, ["html"])
    if ok:
        ok = "<h1" in r.get_json()["html"]
        msg = "渲染出 h1"
    record("功能", "markdown_preview 预览", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "markdown_preview 预览", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.22 文本去重
try:
    data = [("text", "a\nb\na\nc"), ("trim", "1"), ("sort", "1")]
    r = post_tool("text_dedup", data)
    ok, msg = _assert_json_ok(r, ["unique_count"])
    if ok:
        ok = r.get_json()["unique_count"] == 3
        msg = f"去重后 {r.get_json()['unique_count']} 行"
    record("功能", "text_dedup 去重", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "text_dedup 去重", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.23 单位换算
try:
    data = [("category", "length"), ("value", "1"), ("from_unit", "m"), ("to_unit", "km")]
    r = post_tool("unit_convert", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = abs(float(r.get_json()["result"]) - 0.001) < 1e-9
        msg = f"1m = {r.get_json()['result']} km"
    record("功能", "unit_convert 长度换算", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "unit_convert 长度换算", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.24 颜色转换
try:
    data = [("color", "#ff0000")]
    r = post_tool("color_convert", data)
    ok, msg = _assert_json_ok(r, ["rgb"])
    if ok:
        ok = "rgb(255, 0, 0)" == r.get_json()["rgb"]
        msg = r.get_json()["rgb"]
    record("功能", "color_convert 颜色转换", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "color_convert 颜色转换", False, str(exc).splitlines()[0], traceback.format_exc())

# 5.25 数据格式转换 JSON->CSV 与 CSV->XLSX
try:
    data = [("from", "json"), ("to", "csv"), ("text", '[{"a":1,"b":2},{"a":3,"b":4}]')]
    r = post_tool("data_convert", data)
    ok, msg = _assert_json_ok(r, ["result"])
    if ok:
        ok = "a,b" in r.get_json()["result"]
        msg = "JSON→CSV 成功"
    record("功能", "data_convert JSON→CSV", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "data_convert JSON→CSV", False, str(exc).splitlines()[0], traceback.format_exc())

try:
    csv_text = "name,age\nAlice,30\nBob,25"
    data = [("from", "csv"), ("to", "xlsx"), ("text", csv_text)]
    r = post_tool("data_convert", data)
    ok, msg = _assert_json_ok(r, ["url", "mime"])
    record("功能", "data_convert CSV→XLSX", ok, msg)
except Exception as exc:  # noqa: BLE001
    record("功能", "data_convert CSV→XLSX", False, str(exc).splitlines()[0], traceback.format_exc())


# ---------------------------------------------------------------------------
# 6. 鉴权流程（注册 / 登录 / 登出）
# ---------------------------------------------------------------------------
# 注销当前管理员，确保从干净状态开始
try:
    client.get("/logout", follow_redirects=False)
except Exception:  # noqa: BLE001
    pass

# 注册新用户
try:
    r = client.post("/register", data={
        "email": "user@test.com",
        "password": "Passw0rd1",
        "confirm": "Passw0rd1",
        "remember": "y",
    }, follow_redirects=False)
    ok = r.status_code in (301, 302)
    record("鉴权", "注册新用户", ok, f"status={r.status_code}, location={r.headers.get('Location','')}")
except Exception as exc:  # noqa: BLE001
    record("鉴权", "注册新用户", False, str(exc).splitlines()[0], traceback.format_exc())

# 注册重复邮箱应被拒绝
try:
    r = client.post("/register", data={
        "email": "user@test.com",
        "password": "Passw0rd1",
        "confirm": "Passw0rd1",
        "remember": "y",
    }, follow_redirects=False)
    # 应被重定向回登录页（flash 提示）
    ok = r.status_code in (301, 302)
    record("鉴权", "重复邮箱注册被拦截", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("鉴权", "重复邮箱注册被拦截", False, str(exc).splitlines()[0], traceback.format_exc())

# 错误密码登录失败（先登出，确保从匿名状态测试）
try:
    client.get("/logout", follow_redirects=False)
    r = client.post("/login", data={
        "email": "user@test.com",
        "password": "wrongpass",
        "remember": "y",
    }, follow_redirects=False)
    ok = r.status_code in (200, 401)  # 200 重新渲染登录页 或 401
    record("鉴权", "错误密码登录失败", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("鉴权", "错误密码登录失败", False, str(exc).splitlines()[0], traceback.format_exc())

# 正确登录普通用户
try:
    r = client.post("/login", data={
        "email": "user@test.com",
        "password": "Passw0rd1",
        "remember": "y",
    }, follow_redirects=False)
    ok = r.status_code in (301, 302)
    record("鉴权", "普通用户登录", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("鉴权", "普通用户登录", False, str(exc).splitlines()[0], traceback.format_exc())

# 登出
try:
    r = client.get("/logout", follow_redirects=False)
    ok = r.status_code in (301, 302)
    record("鉴权", "登出", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("鉴权", "登出", False, str(exc).splitlines()[0], traceback.format_exc())


# ---------------------------------------------------------------------------
# 7. 管理后台
# ---------------------------------------------------------------------------
# 普通用户访问后台应被拒绝
try:
    client.post("/login", data={
        "email": "user@test.com",
        "password": "Passw0rd1",
        "remember": "y",
    }, follow_redirects=True)
    r = client.get("/admin/")
    ok = r.status_code in (403, 302)  # 403 或跳转登录
    record("后台", "普通用户访问后台被拒绝", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("后台", "普通用户访问后台被拒绝", False, str(exc).splitlines()[0], traceback.format_exc())

# 切回管理员
login_admin()

ADMIN_PAGES = [
    ("/admin/", "仪表盘"),
    ("/admin/users", "用户列表"),
    ("/admin/tools", "工具列表"),
    ("/admin/logs", "日志"),
    ("/admin/settings", "设置"),
]
for url, name in ADMIN_PAGES:
    try:
        r = client.get(url)
        ok = r.status_code == 200
        record("后台", f"GET {url} ({name})", ok, f"status={r.status_code}")
    except Exception as exc:  # noqa: BLE001
        record("后台", f"GET {url} ({name})", False, str(exc).splitlines()[0], traceback.format_exc())

# 保存设置
try:
    r = client.post("/admin/settings", data={
        "site_name": "测试工具箱",
        "site_tagline": "测试标语",
        "daily_free_limit": "10",
        "anon_free_limit": "3",
        "AI_PROVIDER": "mock",
        "AI_BASE_URL": "https://image.pollinations.ai",
        "AI_MODEL": "",
        "AI_API_KEY": "",
    }, follow_redirects=False)
    ok = r.status_code in (301, 302)
    record("后台", "保存设置", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("后台", "保存设置", False, str(exc).splitlines()[0], traceback.format_exc())

# 切换一个工具启用状态
try:
    r = client.post("/admin/tools/pdf_merge/toggle", follow_redirects=False)
    ok = r.status_code in (301, 302)
    record("后台", "切换工具启用状态", ok, f"status={r.status_code}")
    # 切回去，保证环境一致
    client.post("/admin/tools/pdf_merge/toggle", follow_redirects=False)
except Exception as exc:  # noqa: BLE001
    record("后台", "切换工具启用状态", False, str(exc).splitlines()[0], traceback.format_exc())

# 切换普通用户启用状态
try:
    from models import User
    with app.app_context():
        u = User.query.filter_by(email="user@test.com").first()
        uid = u.id if u else None
    if uid:
        r = client.post(f"/admin/users/{uid}/toggle", follow_redirects=False)
        ok = r.status_code in (301, 302)
        record("后台", "切换用户启用状态", ok, f"status={r.status_code}")
        client.post(f"/admin/users/{uid}/toggle", follow_redirects=False)
    else:
        record("后台", "切换用户启用状态", False, "未找到测试用户")
except Exception as exc:  # noqa: BLE001
    record("后台", "切换用户启用状态", False, str(exc).splitlines()[0], traceback.format_exc())

# 设置用户自定义配额
try:
    if uid:
        r = client.post(f"/admin/users/{uid}/limit", data={"tool_id": "json_formatter", "limit": "50"},
                        follow_redirects=False)
        ok = r.status_code in (301, 302)
        record("后台", "设置用户自定义配额", ok, f"status={r.status_code}")
    else:
        record("后台", "设置用户自定义配额", False, "未找到测试用户")
except Exception as exc:  # noqa: BLE001
    record("后台", "设置用户自定义配额", False, str(exc).splitlines()[0], traceback.format_exc())


# ---------------------------------------------------------------------------
# 8. 匿名试用配额
# ---------------------------------------------------------------------------
try:
    client.get("/logout", follow_redirects=False)
    # 匿名使用 json_formatter 一次（在免费额度内）
    r = post_tool("json_formatter", [("text", '{"x":1}'), ("action", "format")])
    ok = (r.get_json(silent=True) or {}).get("ok") is True
    record("配额", "匿名首次使用工具", ok, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("配额", "匿名首次使用工具", False, str(exc).splitlines()[0], traceback.format_exc())


# ---------------------------------------------------------------------------
# 9. 扩展测试：下载链路 / 错误页 / PN 导入导出 / 匿名配额耗尽 / 非法输入 / 静态资源
# ---------------------------------------------------------------------------
login_admin()


def _dl(tool_id, data, expect_mime=None):
    """发起处理并立刻下载返回的文件，验证 200 与 mimetype。"""
    r = post_tool(tool_id, data)
    j = r.get_json(silent=True) or {}
    if not j.get("ok") or not j.get("url"):
        return False, f"process 未返回 ok/url: {j.get('error', j)}"
    url = j["url"]
    if not url.startswith("/"):
        url = "/" + url
    d = client.get(url)
    if d.status_code != 200:
        return False, f"GET {url} -> {d.status_code}"
    if expect_mime and not d.mimetype.startswith(expect_mime):
        return False, f"mime 不符: {d.mimetype} (期望 {expect_mime})"
    return True, f"{url} -> {d.status_code} {d.mimetype} {len(d.data)}B"


_dl_tests = [
    ("pdf_merge",      [("pdfs", (make_pdf(2), "a.pdf")), ("pdfs", (make_pdf(1), "b.pdf"))], "application/pdf"),
    ("pdf_split",      [("pdf", (make_pdf(3), "s.pdf")), ("ranges", "1-2")], "application/pdf"),
    ("pdf_watermark",  [("pdf", (make_pdf(2), "s.pdf")), ("text", "WM")], "application/pdf"),
    ("pdf_rotate",     [("pdf", (make_pdf(2), "s.pdf")), ("angle", "90"), ("scope", "all")], "application/pdf"),
    ("pdf_to_word",    [("pdf", (make_pdf(2), "s.pdf"))], "application/vnd.openxmlformats"),
    ("image_compress", [("image", (make_png(), "p.png")), ("quality", "60")], "image/"),
    ("image_convert",  [("file", (make_png(), "p.png")), ("format", "jpg"), ("quality", "85")], "image/"),
    ("image_to_pdf",   [("files", (make_png(), "p.png"))], "application/pdf"),
    ("qr_gen",         [("text", "hi"), ("size", "10")], "image/"),
    ("word_to_pdf",    [("file", (make_docx(), "d.docx"))], "application/pdf"),
    ("data_convert",   [("from", "csv"), ("to", "xlsx"), ("text", "a,b\n1,2")], "application/vnd.openxmlformats"),
    ("fcst_merge",     [("files", (make_fcst_xlsx(), "f.xlsx"))], "application/vnd.openxmlformats"),
]
for tid, data, mime in _dl_tests:
    try:
        ok, msg = _dl(tid, data, mime)
        record("下载", f"{tid} 下载产物", ok, msg)
    except Exception as exc:  # noqa: BLE001
        record("下载", f"{tid} 下载产物", False, str(exc).splitlines()[0], traceback.format_exc())

# ai_image 下载（status 返回 url）
try:
    r = client.post("/tools/ai-image/generate", data={"prompt": "x", "size": "256x256"},
                    content_type="multipart/form-data", headers=AJ)
    j = r.get_json() or {}
    if j.get("task_id"):
        s = client.get(f"/tools/ai-image/status/{j['task_id']}", headers=AJ).get_json() or {}
        url = s.get("url", "")
        if url:
            d = client.get(url)
            ok = d.status_code == 200 and d.mimetype.startswith("image/")
            record("下载", "ai_image 下载产物", ok, f"{url} -> {d.status_code} {d.mimetype}")
        else:
            record("下载", "ai_image 下载产物", False, f"status 未返回 url: {s}")
    else:
        record("下载", "ai_image 下载产物", False, f"未返回 task_id: {j}")
except Exception as exc:  # noqa: BLE001
    record("下载", "ai_image 下载产物", False, str(exc).splitlines()[0], traceback.format_exc())

# base64 解码图片后下载
try:
    import base64 as _b
    png_bytes = make_png().getvalue()
    b64 = _b.b64encode(png_bytes).decode()
    r = post_tool("base64", [("action", "decode"), ("text", f"data:image/png;base64,{b64}")])
    j = r.get_json() or {}
    url = j.get("url", "")
    if url:
        d = client.get(url)
        ok = d.status_code == 200 and d.mimetype.startswith("image/")
        record("下载", "base64 解码图片下载", ok, f"{url} -> {d.status_code} {d.mimetype}")
    else:
        record("下载", "base64 解码图片下载", False, f"未返回 url: {j}")
except Exception as exc:  # noqa: BLE001
    record("下载", "base64 解码图片下载", False, str(exc).splitlines()[0], traceback.format_exc())

# 错误页
try:
    r = client.get("/this-page-does-not-exist")
    record("边界", "404 错误页", r.status_code == 404, f"status={r.status_code}")
except Exception as exc:  # noqa: BLE001
    record("边界", "404 错误页", False, str(exc).splitlines()[0])

# 静态资源
for sfile in ["/static/js/result.js", "/static/js/tools.js", "/static/js/main.js", "/static/css/style.css"]:
    try:
        r = client.get(sfile)
        record("静态", f"GET {sfile}", r.status_code == 200, f"status={r.status_code}")
    except Exception as exc:  # noqa: BLE001
        record("静态", f"GET {sfile}", False, str(exc).splitlines()[0])

# PN 批量导入
try:
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active
    ws.append(["品号", "原厂料号", "品牌"])
    ws.append(["IMP1", "M-IMP1", "BrandA"])
    ws.append(["IMP2", "M-IMP2", "BrandB"])
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    r = client.post("/tools/fcst-merge/api/pn/import",
                    data={"file": (buf, "pn.xlsx")}, content_type="multipart/form-data", headers=AJ)
    j = r.get_json() or {}
    ok = r.status_code == 200 and j.get("ok") and j.get("added", 0) >= 1
    record("功能", "fcst_merge PN 批量导入", ok, f"status={r.status_code} added={j.get('added')}")
except Exception as exc:  # noqa: BLE001
    record("功能", "fcst_merge PN 批量导入", False, str(exc).splitlines()[0], traceback.format_exc())

# PN 导出
try:
    r = client.get("/tools/fcst-merge/api/pn/export", headers=AJ)
    ok = r.status_code == 200 and "spreadsheet" in (r.mimetype or "")
    record("功能", "fcst_merge PN 导出 xlsx", ok, f"status={r.status_code} {r.mimetype}")
except Exception as exc:  # noqa: BLE001
    record("功能", "fcst_merge PN 导出 xlsx", False, str(exc).splitlines()[0])

# PN 统计
try:
    r = client.get("/tools/fcst-merge/api/pn/stats", headers=AJ)
    j = r.get_json() or {}
    ok = r.status_code == 200 and j.get("ok") and j.get("total", 0) >= 0
    record("功能", "fcst_merge PN 统计", ok, f"status={r.status_code} total={j.get('total')}")
except Exception as exc:  # noqa: BLE001
    record("功能", "fcst_merge PN 统计", False, str(exc).splitlines()[0])

# 匿名配额耗尽：新客户端连续调用同一工具，第 4 次应被拦截
try:
    c2 = app.test_client()
    last_status = None
    blocked_correctly = False
    for i in range(4):
        rr = c2.post("/tools/json-formatter/process",
                     data={"text": '{"x":1}', "action": "format"},
                     content_type="multipart/form-data", headers=AJ)
        last_status = rr.status_code
        jj = rr.get_json(silent=True) or {}
        if i < 3:
            if not jj.get("ok"):
                blocked_correctly = False
                break
        else:
            # 第 4 次应被拦截（429 或非 ok）
            blocked_correctly = (rr.status_code == 429 or not jj.get("ok"))
    record("配额", "匿名配额耗尽拦截", blocked_correctly, f"第4次 status={last_status}")
except Exception as exc:  # noqa: BLE001
    record("配额", "匿名配额耗尽拦截", False, str(exc).splitlines()[0], traceback.format_exc())

# 非法输入处理
_invalid = [
    ("json_formatter", [("text", "{bad json}"), ("action", "format")], 400),
    ("hash_calc",      [("algo", "sha256")], 400),  # 无文本无文件
    ("unit_convert",   [("category", "length"), ("value", "abc"), ("from_unit", "m"), ("to_unit", "km")], 400),
    ("color_convert",  [("color", "zzzzz")], 400),
    ("uuid_gen",       [("count", "abc"), ("version", "4")], 400),
]
for tid, data, expect in _invalid:
    try:
        r = post_tool(tid, data)
        record("边界", f"{tid} 非法输入 -> {expect}", r.status_code == expect, f"status={r.status_code}")
    except Exception as exc:  # noqa: BLE001
        record("边界", f"{tid} 非法输入 -> {expect}", False, str(exc).splitlines()[0])


# ---------------------------------------------------------------------------
# 10. 生成报告
# ---------------------------------------------------------------------------
_write_reports()

# 退出码：有失败则非零
sys.exit(1 if any(not r["ok"] for r in RESULTS) else 0)
